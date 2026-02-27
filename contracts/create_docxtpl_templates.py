"""
Script pour créer des templates Word compatibles avec DocxTemplate
Les templates créés avec python-docx ne fonctionnent pas correctement avec DocxTemplate
car Word fragmente les placeholders Jinja2 en plusieurs runs XML.

Cette version crée des templates minimaux qui seront ensuite édités manuellement.
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_simple_docxtpl_template(entity, contract_type, gender, output_dir):
    """
    Crée un template Word simple compatible DocxTemplate
    """
    doc = Document()
    
    # Marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
    
    # EN-TÊTE ENTREPRISE
    if entity == 'nantes_urgences':
        company_name = 'SARL NANTES URGENCES SANSOUCY'
        address = '8 Rue de Remouleur, 44800 SAINT-HERBLAIN'
        siret = 'SIRET: 48805076600028'
    else:  # ambulances_sansoucy
        company_name = 'SARL AMBULANCES SANSOUCY'
        address = '2 avenue de la Véra Cruz, 44600 SAINT NAZAIRE'
        siret = 'SIRET: 38026793000036'
    
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name)
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph(address)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(siret)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # TITRE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = 'CONTRAT À DURÉE INDÉTERMINÉE' if contract_type == 'cdi' else 'CONTRAT À DURÉE DÉTERMINÉE'
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TEXTE DU CONTRAT - On ajoute tout le texte d'un coup pour éviter la fragmentation
    # Note: Les variables Jinja2 seront ajoutées manuellement après pour éviter la fragmentation XML
    
    doc.add_paragraph('ENTRE:').runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph(f'{company_name}')
    doc.add_paragraph(address)
    doc.add_paragraph()
    
    doc.add_paragraph('ET:').runs[0].bold = True
    doc.add_paragraph()
    
    # Informations employé - PLACEHOLDER À REMPLACER MANUELLEMENT
    civility = 'Monsieur' if gender == 'M' else 'Madame'
    doc.add_paragraph(f'{civility} VAR_FIRST_NAME VAR_LAST_NAME')
    doc.add_paragraph('Adresse: VAR_ADDRESS')
    doc.add_paragraph('VAR_CITY_POSTALCODE')
    doc.add_paragraph('Né(e) le: VAR_BIRTH_DATE à VAR_BIRTH_PLACE (VAR_BIRTH_PLACE_CODE)')
    doc.add_paragraph('N° Sécurité Sociale: VAR_SSN')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Articles du contrat
    doc.add_paragraph('ARTICLE 1: ENGAGEMENT').runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph(f'{civility} VAR_FIRST_NAME VAR_LAST_NAME est engagé(e) en qualité d\'ambulancier à compter du VAR_START_DATE.')
    
    doc.add_paragraph()
    doc.add_paragraph('ARTICLE 2: RÉMUNÉRATION').runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph('Taux horaire: VAR_HOURLY_RATE €/h')
    doc.add_paragraph('Salaire mensuel brut: VAR_MONTHLY_SALARY € (base VAR_MONTHLY_HOURS h/mois)')
    
    if contract_type == 'cdd':
        doc.add_paragraph()
        doc.add_paragraph('ARTICLE 3: DURÉE DU CONTRAT').runs[0].bold = True
        doc.add_paragraph()
        doc.add_paragraph('Date de fin: VAR_END_DATE')
        doc.add_paragraph('Période d\'essai: VAR_TRIAL_PERIOD')
        doc.add_paragraph('Motif du CDD: VAR_CDD_REASON')
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Fait à St-Herblain/St-Nazaire, le VAR_CONTRACT_DATE')
    
    # Sauvegarder
    os.makedirs(output_dir, exist_ok=True)
    gender_suffix = 'homme' if gender == 'M' else 'femme'
    filename = f'template_simple_{entity}_{contract_type}_{gender_suffix}.docx'
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    print(f"✓ Créé: {filename}")
    print(f"  IMPORTANT: Ouvrir ce fichier dans Word et remplacer les VAR_* par les placeholders Jinja2:")
    print(f"  VAR_FIRST_NAME → {{{{ employee_first_name }}}}")
    print(f"  VAR_LAST_NAME → {{{{ employee_last_name }}}}")
    print(f"  etc.")
    
    return filepath


if __name__ == '__main__':
    output_dir = './templates/word_templates'
    
    print("Création de templates simples compatibles DocxTemplate...\n")
    
    for entity in ['nantes_urgences', 'ambulances_sansoucy']:
        for contract_type in ['cdi', 'cdd']:
            for gender in ['M', 'F']:
                create_simple_docxtpl_template(entity, contract_type, gender, output_dir)
    
    print("\n" + "="*70)
    print("IMPORTANT:")
    print("Les templates créés contiennent des placeholders VAR_*")
    print("Vous devez ouvrir chaque fichier dans Microsoft Word et remplacer:")
    print("  VAR_FIRST_NAME par {{employee_first_name}}")
    print("  VAR_LAST_NAME par {{employee_last_name}}")
    print("  VAR_ADDRESS par {{employee_address}}")
    print("  VAR_CITY_POSTALCODE par {{employee_city_postalcode}}")
    print("  VAR_BIRTH_DATE par {{employee_birth_date}}")
    print("  VAR_BIRTH_PLACE par {{employee_birth_place}}")
    print("  VAR_BIRTH_PLACE_CODE par {{employee_birth_place_code}}")
    print("  VAR_SSN par {{employee_social_security_number}}")
    print("  VAR_START_DATE par {{start_date}}")
    print("  VAR_HOURLY_RATE par {{hourly_rate}}")
    print("  VAR_MONTHLY_SALARY par {{monthly_salary}}")
    print("  VAR_MONTHLY_HOURS par {{monthly_hours}}")
    print("  VAR_END_DATE par {{end_date}}")
    print("  VAR_TRIAL_PERIOD par {{trial_period}}")
    print("  VAR_CDD_REASON par {{cdd_reason}}")
    print("  VAR_CONTRACT_DATE par {{contract_date}}")
    print("="*70)
