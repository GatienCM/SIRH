"""
Générateur de templates CDD réalistes basés sur les contrats réels de Nantes Urgences Sansoucy.
Crée 8 variantes: 2 entités × 2 types contrats × 2 genres avec accords grammaticaux.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, fill_color):
    """Define cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def get_gender_agreements(gender):
    """
    Returns grammatical agreements based on gender.
    
    Args:
        gender: 'M' (masculine) or 'F' (feminine)
        
    Returns:
        dict: Dictionary with grammatical variants
    """
    if gender == 'F':
        return {
            'civility': 'Madame',
            'engage': 'engagée',
            'ne': 'née',
            'il_elle': 'elle',
            'son_sa': 'sa',
            'est_was': 'a',
            'se': 'se',
            'place': 'placée',
            'accepte': 'accepte',
        }
    else:
        return {
            'civility': 'Monsieur',
            'engage': 'engagé',
            'ne': 'né',
            'il_elle': 'il',
            'son_sa': 'son',
            'est_was': 'a',
            'se': 's\'',
            'place': 'placé',
            'accepte': 'accepte',
        }


def create_nantes_urgences_cdd_template(contract_type, gender, output_dir):
    """
    Create a realistic CDD contract template for Nantes Urgences Sansoucy.
    
    Args:
        contract_type: 'cdi' or 'cdd'
        gender: 'M' (masculine) or 'F' (feminine)
        output_dir: Output directory for the template file
        
    Returns:
        str: Full path to created template file
    """
    agreements = get_gender_agreements(gender)
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
    
    # HEADER - Company info
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run('NANTES URGENCES SANSOUCY')
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Address
    addr_para = doc.add_paragraph('8 Rue de Remouleur, 44800 SAINT-HERBLAIN')
    addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = addr_para.runs[0]
    run.font.size = Pt(11)
    
    # SIRET and URSSAF
    siret_para = doc.add_paragraph('SIRET: 48805076600028')
    siret_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = siret_para.runs[0]
    run.font.size = Pt(11)
    
    urssaf_para = doc.add_paragraph('URSSAF NANTES 527201905363')
    urssaf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = urssaf_para.runs[0]
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # CONTRACT TITLE
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if contract_type == 'cdd':
        run = title_para.add_run('CONTRAT À DURÉE DÉTERMINÉE')
    else:
        run = title_para.add_run('CONTRAT À DURÉE INDÉTERMINÉE')
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # BETWEEN
    between_para = doc.add_paragraph()
    between_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = between_para.add_run('ENTRE:')
    run.font.size = Pt(11)
    run.font.bold = True
    
    doc.add_paragraph('NANTES URGENCES SANSOUCY', style='Normal')
    doc.add_paragraph('8 Rue de Remouleur, 44800 SAINT-HERBLAIN')
    doc.add_paragraph('Représentée par Monsieur Patrice BORÉ, Direction')
    
    doc.add_paragraph()
    
    doc.add_paragraph('ET')
    doc.add_paragraph()
    
    doc.add_paragraph(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}}')
    doc.add_paragraph('{{{{ employee_address }}}}')
    doc.add_paragraph('{{{{ employee_city_postalcode }}}}')
    doc.add_paragraph(f'{agreements["ne"]} le: {{{{ employee_birth_date }}}} à {{{{ employee_birth_place }}}} ({{{{ employee_birth_place_code }}}})' )
    doc.add_paragraph('Immatriculé: {{ employee_social_security_number }}')
    
    doc.add_paragraph()  # Spacing
    
    # RAPPEL
    p = doc.add_paragraph()
    run = p.add_run('IL A ÉTÉ RAPPELÉ CE QUI SUIT')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # EMPLOYMENT STATUS
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run(f'est {agreements["engage"]} en qualité d\'ambulancier.')
    
    doc.add_paragraph()
    
    # DECLARATION
    p = doc.add_paragraph()
    p.add_run(
        f'La déclaration préalable à l\'embauche de {agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} '
        f'{agreements["est_was"]} effectuée à l\'URSSAF de Nantes et {agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} '
        'pourra exercer auprès de cet organisme son droit d\'accès et de rectification que lui confère la loi 78.17 du 6 janvier 1978.'
    )
    
    doc.add_paragraph()
    
    # ARTICLE 1: ATTRIBUTION AND EMPLOYMENT
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 1: ATTRIBUTION ET EMPLOI')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run(
        f'{agreements["est_was"]} {agreements["engage"]} par la société NANTES URGENCES SANSOUCY en qualité d\'ambulancier, '
        'de la convention collective des transports routiers (IDCC 16), applicable à l\'activité.'
    )
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'En cette qualité {agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run(f'{agreements["il_elle"]} aura pour mission:')
    
    doc.add_paragraph('Le transport de personnes en ambulance et VSL', style='List Bullet')
    doc.add_paragraph('Et plus généralement le transport de personnes.', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run(
        f'sera {agreements["place"]} sous l\'autorité hiérarchique de la direction de la société NANTES URGENCES SANSOUCY, '
        f'{agreements["il_elle"]} devra également respecter les missions et les directives qui {agreements["se"]}lui seront données par les régulateurs.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ARTICLE 2: WORKING CONDITIONS
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 2: CONDITIONS GÉNÉRALES DE TRAVAIL')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 2.1 Work location
    p = doc.add_paragraph()
    run = p.add_run('2.1 Lieu de travail')
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run('exercera ses fonctions au sein des établissements de la société NANTES URGENCES SANSOUCY basée à St-Herblain ou Carquefou.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Toutefois, ').font.size = Pt(11)
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run(
        f'{agreements["accepte"]} par avance qu\'en fonction des nécessités de l\'entreprise, {agreements["il_elle"]} soit amené à changer de lieu de travail, '
        'et ce, dans les zones géographiques où la société exerce ou exercera son activité.'
    ).font.size = Pt(11)
    
    # 2.2 Paid leave
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('2.2 Congés payés')
    run.font.bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Conformément aux dispositions légales en vigueur et celles de la convention collective applicable, '
        'le salarié bénéficiera de 2.08 jours ouvrés de congés par mois de travail effectif, '
        'acquis sur la période courant du 1er juin au 31 mai de l\'année suivante.'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph('La direction fixe les périodes de congés, en fonction des nécessités du service, après concertation du personnel.')
    
    # 2.3 Retirement
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('2.3 Caisse de retraite complémentaire')
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run('sera affilié à:')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Pour la retraite: UGRR - BP 501 - 75421 PARIS CEDEX 09')
    doc.add_paragraph('Pour la prévoyance: ALLIANZ PRÉVOYANCE - 1 cours Michelet - CS 30051 92076 PARIS LA DEFENSE CEDEX')
    
    # 2.4 General conditions
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('2.4 Conditions générales d\'exercice des fonctions')
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run(
        f'{agreements["se"]}engage à se conformer strictement aux instructions de la direction concernant les conditions d\'exécution du travail, '
        'et à respecter les plannings journaliers qui seront établis par les régulateurs.'
    )
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run(f'déclare ne faire l\'objet d\'aucune restriction administrative ou judiciaire quant à l\'utilisation de {agreements["son_sa"]} permis de conduire, '
        'toutes catégories confondues.')
    
    # ARTICLE 3: DISCIPLINE AND SECURITY
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 3: DISCIPLINE ET SÉCURITÉ')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run('déclare que toutes les coordonnées figurant à l\'entête des présentes sont exactes, et s\'oblige à prévenir l\'employeur de toutes modifications les affectant.')
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Il s\'oblige à prévenir sans délai la société NANTES URGENCES SANSOUCY de toute absence quelle qu\'en soit la cause '
        'et de la justifier dans les 48h. Sinon cela correspondra à une absence injustifiée, qui, répétée, peut engendrer une sanction disciplinaire.'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph('En cas de maladie, il faut faire parvenir un arrêt de travail dans les 48 heures de l\'arrêt au service ressources humaines de l\'entreprise.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run('reconnaît avoir pris connaissance du règlement intérieur en vigueur dans l\'établissement. '
        'Tout manquement au présent règlement pourra donner lieu à des poursuites disciplinaires et à un éventuel licenciement pour faute.')
    
    # ARTICLE 4: CONFIDENTIALITY
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 4: CONFIDENTIALITÉ ET SECRET PROFESSIONNEL')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        f'Compte tenu des fonctions confiées à {agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }}, '
        'celui-ci est tenu par un secret professionnel tant en ce qui concerne l\'identité des clients transportés que leur destination.'
    )
    
    # ARTICLE 5: SALARY
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 5: SALAIRE')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Le salaire de ').font.size = Pt(11)
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }}').font.bold = True
    p.add_run(', SMPG (Salaire Minimum Professionnel Garanti), se décompose comme suit:').font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Salary table
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Emploi'
    header_cells[1].text = 'Taux'
    header_cells[2].text = 'À compter du 1er août 2025 (base {{ monthly_hours }}h / mois)'
    
    for cell in header_cells:
        set_cell_background(cell, 'D9D9D9')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data row
    data_cells = table.rows[1].cells
    data_cells[0].text = 'DEA'
    data_cells[1].text = '{{ hourly_rate }}€'
    data_cells[2].text = '{{ monthly_salary }}€'
    
    doc.add_paragraph()
    
    # Additional allowances
    p = doc.add_paragraph()
    p.add_run('En sus du SMPG, il pourra être versé à ').font.size = Pt(11)
    p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
    p.add_run('les indemnités suivantes, dans les termes de l\'accord du 04 mai 2000:').font.size = Pt(11)
    
    doc.add_paragraph('IDAJ (Indemnité Dépassement d\'amplitude Journalière)')
    doc.add_paragraph(
        'Tâches complémentaires: elles se trouvent définies à l\'accord cadre et donneront lieu à un paiement spécifique '
        'chaque fois qu\'elles auront été effectuées.'
    )
    
    # CDD-specific contract details
    if contract_type == 'cdd':
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('ARTICLE 6: DURÉE ET CAUSE DU CONTRAT')
        run.font.bold = True
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run(f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} ').font.bold = True
        p.add_run('est engagé à titre précaire du {{ start_date }} au {{ end_date }}, avec période d\'essai de {{ trial_period }}.')
        
        p = doc.add_paragraph()
        p.add_run('Motif du recours au CDD: {{ cdd_reason }}')
        
        # ARTICLE 7: INDEMNITÉ DE FIN DE CONTRAT
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('ARTICLE 7: INDEMNITÉ DE FIN DE CONTRAT')
        run.font.bold = True
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        doc.add_paragraph(
            'Lorsque, à l\'issue de ce contrat, les relations de travail ne se poursuivent pas par un contrat à durée indéterminée, '
            f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }} a droit à une indemnité de fin de contrat égale à 10% '
            'de la rémunération totale brute versée.'
        )
    
    # CONTRACT DATES AND SIGNATURE
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Fait à St-Herblain, le {{ contract_date }}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature table
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = 'Table Grid'
    
    # Employee signature
    emp_cells = sig_table.rows[0].cells
    emp_cells[0].text = f'{agreements["civility"]} {{ employee_first_name }} {{ employee_last_name }}:'
    emp_para = emp_cells[0].paragraphs[0]
    emp_para.runs[0].font.bold = True
    
    emp_desc = emp_cells[0].add_paragraph('Bon pour accord, lu et approuvé')
    
    # Company signature
    comp_cells = sig_table.rows[0].cells
    comp_cells[1].text = 'NANTES URGENCES SANSOUCY'
    comp_para = comp_cells[1].paragraphs[0]
    comp_para.runs[0].font.bold = True
    comp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    comp_desc = comp_cells[1].add_paragraph('Monsieur Patrice BORÉ')
    comp_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    comp_desc2 = comp_cells[1].add_paragraph('Direction')
    comp_desc2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    comp_desc2.runs[0].font.bold = True
    
    # Save template
    os.makedirs(output_dir, exist_ok=True)
    gender_suffix = 'homme' if gender == 'M' else 'femme'
    filename = f'contrat_nantes_urgences_{contract_type}_{gender_suffix}.docx'
    filepath = os.path.join(output_dir, filename)
    
    doc.save(filepath)
    return filepath


def create_ambulances_sansoucy_cdd_template(contract_type, gender, output_dir):
    """
    Create a realistic CDD contract template for Ambulances Sansoucy.
    
    Args:
        contract_type: 'cdi' or 'cdd'
        gender: 'M' (masculine) or 'F' (feminine)
        output_dir: Output directory for the template file
        
    Returns:
        str: Full path to created template file
    """
    agreements = get_gender_agreements(gender)
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
    
    # HEADER - Company info
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run('SARL AMBULANCES SANSOUCY')
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Address
    addr_para = doc.add_paragraph('2 avenue de la Véra Cruz, 44600 SAINT NAZAIRE')
    addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = addr_para.runs[0]
    run.font.size = Pt(11)
    
    # SIRET and URSSAF
    siret_para = doc.add_paragraph('SIRET: 38026793000036')
    siret_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = siret_para.runs[0]
    run.font.size = Pt(11)
    
    urssaf_para = doc.add_paragraph('URSSAF NANTES 527201905363')
    urssaf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = urssaf_para.runs[0]
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # CONTRACT TITLE
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if contract_type == 'cdd':
        run = title_para.add_run('CONTRAT À DURÉE DÉTERMINÉE')
    else:
        run = title_para.add_run('CONTRAT À DURÉE INDÉTERMINÉE')
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # BETWEEN
    between_para = doc.add_paragraph()
    between_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = between_para.add_run('ENTRE:')
    run.font.size = Pt(11)
    run.font.bold = True
    
    doc.add_paragraph('SARL AMBULANCES SANSOUCY', style='Normal')
    doc.add_paragraph('2 avenue de la Véra Cruz, 44600 SAINT NAZAIRE')
    doc.add_paragraph('Représentée par Monsieur Bruno SANSOUCY, Gérant')
    
    doc.add_paragraph()
    doc.add_paragraph('ET')
    doc.add_paragraph()
    
    doc.add_paragraph(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}}')
    doc.add_paragraph('{{{{ employee_address }}}}')
    doc.add_paragraph('{{{{ employee_city_postalcode }}}}')
    doc.add_paragraph(f'{agreements["ne"]} le: {{{{ employee_birth_date }}}} à {{{{ employee_birth_place }}}} ({{{{ employee_birth_place_code }}}})' )
    doc.add_paragraph('Immatriculé: {{{{ employee_social_security_number }}}}')
    
    doc.add_paragraph()  # Spacing
    
    # RAPPEL
    p = doc.add_paragraph()
    run = p.add_run('IL A ÉTÉ RAPPELÉ CE QUI SUIT')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # EMPLOYMENT STATUS
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run(f'est {agreements["engage"]} en qualité d\'ambulancier.')
    
    doc.add_paragraph()
    
    # DECLARATION
    p = doc.add_paragraph()
    p.add_run(
        f'La déclaration préalable à l\'embauche de {agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} '
        f'{agreements["est_was"]} effectuée à l\'URSSAF de Nantes et {agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} '
        'pourra exercer auprès de cet organisme son droit d\'accès et de rectification que lui confère la loi 78.17 du 6 janvier 1978.'
    )
    
    doc.add_paragraph()
    
    # ARTICLE 1: ATTRIBUTION AND EMPLOYMENT
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 1: ATTRIBUTION ET EMPLOI')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run(
        f'{agreements["est_was"]} {agreements["engage"]} par la société SARL AMBULANCES SANSOUCY en qualité d\'ambulancier, '
        'de la convention collective des transports routiers (IDCC 16), applicable à l\'activité.'
    )
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'En cette qualité {agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run(f'{agreements["il_elle"]} aura pour mission:')
    
    doc.add_paragraph('Le transport de personnes en ambulance et VSL', style='List Bullet')
    doc.add_paragraph('Et plus généralement le transport de personnes.', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run(
        f'sera {agreements["place"]} sous l\'autorité hiérarchique des cogérants de la société SARL AMBULANCES SANSOUCY, '
        f'{agreements["il_elle"]} devra également respecter les plannings et les directives qui {agreements["se"]}lui seront données par les régulateurs.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ARTICLE 2: WORKING CONDITIONS
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 2: CONDITIONS GÉNÉRALES DE TRAVAIL')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 2.1 Work location
    p = doc.add_paragraph()
    run = p.add_run('2.1 Lieu de travail')
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run('exercera ses fonctions à partir du siège de la société SARL AMBULANCES SANSOUCY.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Toutefois, ').font.size = Pt(11)
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run(
        f'{agreements["accepte"]} par avance qu\'en fonction des nécessités de l\'entreprise, {agreements["il_elle"]} soit amené à changer de lieu de travail, '
        'et ce, dans les zones géographiques où la société exerce ou exercera son activité.'
    ).font.size = Pt(11)
    
    # 2.2 Paid leave
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('2.2 Congés payés')
    run.font.bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Conformément aux dispositions légales en vigueur et celles de la convention collective applicable, '
        'le salarié bénéficiera de 2.08 jours ouvrés de congés par mois de travail effectif, '
        'acquis sur la période courant du 1er juin au 31 mai de l\'année suivante.'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph('La direction fixe les périodes de congés, en fonction des nécessités du service, après concertation du personnel.')
    
    # 2.3 Retirement
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('2.3 Affiliations: Retraite, Santé et Prévoyance')
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run('sera affilié à:')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Pour la retraite: UGRR - BP 501 - 75421 PARIS CEDEX 09')
    doc.add_paragraph('Pour la complémentaire santé: HARMONIE MUTUELLE - 44824 ST HERBLAIN CEDEX')
    doc.add_paragraph('Pour la prévoyance: ALLIANZ VIE - 1 Cours Michelet - 92076 PARIS LA DEFENSE CEDEX')
    
    # 2.4 General conditions
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('2.4 Conditions générales d\'exercice des fonctions')
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run(
        f'{agreements["se"]}engage à se conformer strictement aux instructions de la direction concernant les conditions d\'exécution du travail, '
        'et à respecter les plannings qui seront établis par les régulateurs.'
    )
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run(f'déclare ne faire l\'objet d\'aucune restriction administrative ou judiciaire quant à l\'utilisation de {agreements["son_sa"]} permis de conduire, '
        'toutes catégories confondues.')
    
    # ARTICLE 3: DISCIPLINE AND SECURITY
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 3: DISCIPLINE ET SÉCURITÉ')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run('déclare que toutes les coordonnées figurant à l\'entête des présentes sont exactes, et s\'oblige à prévenir l\'employeur de toutes modifications les affectant.')
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        f'{agreements["se"].capitalize()}oblige à prévenir sans délai la société SARL AMBULANCES SANSOUCY de toute absence '
        'et, en cas de maladie, faire parvenir un arrêt de travail dans les 48 heures.'
    )
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run('reconnaît avoir pris connaissance du règlement intérieur en vigueur dans l\'établissement. '
        'Tout manquement au présent règlement pourra donner lieu à des poursuites disciplinaires et à un éventuel licenciement pour faute.')
    
    # ARTICLE 4: CONFIDENTIALITY
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 4: CONFIDENTIALITÉ ET SECRET PROFESSIONNEL')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        f'Compte tenu des fonctions confiées à {agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}}, '
        'celui-ci est tenu par un secret professionnel tant en ce qui concerne l\'identité des clients transportés que leur destination.'
    )
    
    # ARTICLE 5: SALARY
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('ARTICLE 5: SALAIRE')
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Le salaire de ').font.size = Pt(11)
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}}').font.bold = True
    p.add_run(', SMPG (Salaire Minimum Professionnel Garanti), se décompose comme suit:').font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Salary table (Ambulances Sansoucy uses AA level instead of DEA)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Emploi'
    header_cells[1].text = 'Taux'
    header_cells[2].text = 'À compter du 1er janvier 2025 (base {{{{ monthly_hours }}}}h / mois)'
    
    for cell in header_cells:
        set_cell_background(cell, 'D9D9D9')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data row
    data_cells = table.rows[1].cells
    data_cells[0].text = 'AA'
    data_cells[1].text = '{{{{ hourly_rate }}}}€'
    data_cells[2].text = '{{{{ monthly_salary }}}}€'
    
    doc.add_paragraph()
    
    # Additional allowances
    p = doc.add_paragraph()
    p.add_run('En sus du SMPG, il pourra être versé à ').font.size = Pt(11)
    p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
    p.add_run('les indemnités suivantes, dans les termes de l\'accord du 04 mai 2000:').font.size = Pt(11)
    
    doc.add_paragraph('IDAJ (Indemnité Dépassement d\'amplitude Journalière)')
    doc.add_paragraph(
        'Tâches complémentaires: elles se trouvent définies à l\'accord cadre et donneront lieu à un paiement spécifique '
        'chaque fois qu\'elles auront été effectuées.'
    )
    
    # CDD-specific contract details
    if contract_type == 'cdd':
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('ARTICLE 6: DURÉE ET CAUSE DU CONTRAT')
        run.font.bold = True
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run(f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} ').font.bold = True
        p.add_run('est engagé à titre précaire du {{{{ start_date }}}} au {{{{ end_date }}}}, avec période d\'essai de {{{{ trial_period }}}}.')
        
        p = doc.add_paragraph()
        p.add_run('Motif du recours au CDD: {{{{ cdd_reason }}}}')
        
        # ARTICLE 7: INDEMNITÉ DE FIN DE CONTRAT
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('ARTICLE 7: INDEMNITÉ DE FIN DE CONTRAT')
        run.font.bold = True
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        doc.add_paragraph(
            'Lorsque, à l\'issue de ce contrat, les relations de travail ne se poursuivent pas par un contrat à durée indéterminée, '
            f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}} a droit à une indemnité de fin de contrat égale à 10% '
            'de la rémunération totale brute versée.'
        )
    
    # CONTRACT DATES AND SIGNATURE
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Fait à St-Nazaire, le {{{{ contract_date }}}}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature table
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = 'Table Grid'
    
    # Employee signature
    emp_cells = sig_table.rows[0].cells
    emp_cells[0].text = f'{agreements["civility"]} {{{{ employee_first_name }}}} {{{{ employee_last_name }}}}:'
    emp_para = emp_cells[0].paragraphs[0]
    emp_para.runs[0].font.bold = True
    
    emp_desc = emp_cells[0].add_paragraph('Bon pour accord, lu et approuvé')
    
    # Company signature
    comp_cells = sig_table.rows[0].cells
    comp_cells[1].text = 'SARL AMBULANCES SANSOUCY'
    comp_para = comp_cells[1].paragraphs[0]
    comp_para.runs[0].font.bold = True
    comp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    comp_desc = comp_cells[1].add_paragraph('Monsieur Bruno SANSOUCY')
    comp_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    comp_desc2 = comp_cells[1].add_paragraph('Gérant')
    comp_desc2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    comp_desc2.runs[0].font.bold = True
    
    # Save template
    os.makedirs(output_dir, exist_ok=True)
    gender_suffix = 'homme' if gender == 'M' else 'femme'
    filename = f'contrat_ambulances_sansoucy_{contract_type}_{gender_suffix}.docx'
    filepath = os.path.join(output_dir, filename)
    
    doc.save(filepath)
    return filepath


if __name__ == '__main__':
    # Test - create all 8 template variants
    output_dir = '/tmp/contract_templates'
    os.makedirs(output_dir, exist_ok=True)
    
    # Nantes Urgences Sansoucy templates
    for contract_type in ['cdi', 'cdd']:
        for gender in ['M', 'F']:
            filepath = create_nantes_urgences_cdd_template(contract_type, gender, output_dir)
            print(f'✓ Created: {filepath}')
    
    # Ambulances Sansoucy templates
    for contract_type in ['cdi', 'cdd']:
        for gender in ['M', 'F']:
            filepath = create_ambulances_sansoucy_cdd_template(contract_type, gender, output_dir)
            print(f'✓ Created: {filepath}')
