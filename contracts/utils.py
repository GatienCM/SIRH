"""
Utilitaires pour la génération de contrats de travail
"""
import os
from datetime import datetime
from decimal import Decimal
from django.conf import settings
from django.core.files.base import ContentFile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate


def create_contract_template():
    """
    Crée un template de contrat de travail au format Word compatible avec docxtpl
    Cette fonction génère le template s'il n'existe pas encore
    """
    template_dir = os.path.join(settings.BASE_DIR, 'templates', 'word_templates')
    template_path = os.path.join(template_dir, 'contrat_travail_template.docx')
    
    # Si le template existe déjà, ne pas le recréer
    if os.path.exists(template_path):
        return template_path
    
    # Créer le dossier si nécessaire
    os.makedirs(template_dir, exist_ok=True)
    
    # Créer le document
    doc = Document()
    
    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CONTRAT DE TRAVAIL')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # Sous-titre avec variables Jinja2
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('{{ contract_type_display }}')
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()
    
    # ENTRE LES SOUSSIGNÉS
    doc.add_paragraph('ENTRE LES SOUSSIGNÉS :').runs[0].bold = True
    doc.add_paragraph()
    
    # Employeur
    p = doc.add_paragraph()
    p.add_run('L\'EMPLOYEUR\n').bold = True
    p.add_run('Nom de l\'entreprise : [Nom de votre entreprise]\n')
    p.add_run('Adresse : [Adresse de l\'entreprise]\n')
    p.add_run('SIRET : [Numéro SIRET]\n')
    p.add_run('Représentée par : [Nom du représentant légal]\n')
    
    doc.add_paragraph()
    doc.add_paragraph('D\'UNE PART,').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Salarié avec variables Jinja2
    p = doc.add_paragraph()
    p.add_run('ET LE SALARIÉ\n').bold = True
    p.add_run('Nom : {{ employee_last_name }}\n')
    p.add_run('Prénom : {{ employee_first_name }}\n')
    p.add_run('Date de naissance : {{ employee_birth_date }}\n')
    p.add_run('Lieu de naissance : {{ employee_birth_place }}\n')
    p.add_run('N° Sécurité sociale : {{ employee_social_security }}\n')
    p.add_run('Adresse : {{ employee_address }}\n')
    
    doc.add_paragraph()
    doc.add_paragraph('D\'AUTRE PART,').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()
    
    # ARTICLE 1 - ENGAGEMENT
    doc.add_paragraph('ARTICLE 1 - ENGAGEMENT').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Le salarié est engagé à compter du {{ start_date }} en qualité de {{ employee_profession }}.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Type de contrat : {{ contract_type_display }}')
    
    # Pour les CDD - utilisation de blocs Jinja2
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('{% if end_date %}Date de fin de contrat : {{ end_date }}{% endif %}')
    
    doc.add_paragraph()
    
    # ARTICLE 2 - PÉRIODE D'ESSAI
    doc.add_paragraph('ARTICLE 2 - PÉRIODE D\'ESSAI').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('{% if trial_end_date %}')
    p.add_run('Le contrat est conclu avec une période d\'essai qui prendra fin le {{ trial_end_date }}.')
    p.add_run('{% else %}')
    p.add_run('Le contrat est conclu sans période d\'essai.')
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    
    # ARTICLE 3 - DURÉE DU TRAVAIL
    doc.add_paragraph('ARTICLE 3 - DURÉE DU TRAVAIL').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('La durée hebdomadaire de travail est fixée à {{ working_hours_per_week }} heures par semaine.')
    
    doc.add_paragraph()
    
    # ARTICLE 4 - RÉMUNÉRATION
    doc.add_paragraph('ARTICLE 4 - RÉMUNÉRATION').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('{% if monthly_salary %}')
    p.add_run('Le salarié percevra une rémunération mensuelle brute de {{ monthly_salary }} euros.')
    p.add_run('{% endif %}')
    
    p = doc.add_paragraph()
    p.add_run('{% if hourly_rate %}')
    p.add_run('Taux horaire brut : {{ hourly_rate }} euros.')
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('La rémunération sera versée mensuellement.')
    
    doc.add_paragraph()
    
    # ARTICLE 5 - CONVENTION COLLECTIVE
    doc.add_paragraph('ARTICLE 5 - CONVENTION COLLECTIVE').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Le présent contrat est soumis à la {{ collective_agreement }}.')
    
    doc.add_paragraph()
    
    # ARTICLE 6 - MÉDECINE DU TRAVAIL
    doc.add_paragraph('ARTICLE 6 - MÉDECINE DU TRAVAIL').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('{% if occupational_health_service %}')
    p.add_run('Service de médecine du travail : {{ occupational_health_service }}')
    p.add_run('{% else %}')
    p.add_run('Le salarié sera convoqué à une visite médicale d\'embauche conformément à la réglementation.')
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()
    
    # Signatures
    doc.add_paragraph('Fait en double exemplaire, le {{ signature_date }}')
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table pour les signatures
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # En-têtes
    table.cell(0, 0).text = 'L\'EMPLOYEUR'
    table.cell(0, 1).text = 'LE SALARIÉ'
    
    # Lignes vides pour les signatures
    table.cell(1, 0).text = '\n\n\n'
    table.cell(1, 1).text = '\n\n\n'
    
    # Noms
    table.cell(2, 0).text = '[Nom et signature]'
    table.cell(2, 1).text = '{{ employee_first_name }} {{ employee_last_name }}\nSignature :'
    
    # Centrer le texte des signatures
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sauvegarder
    doc.save(template_path)
    
    return template_path


def create_entity_template(entity_name):
    """
    Crée un template de contrat personnalisé pour une entité spécifique
    
    Args:
        entity_name: 'nantes_urgences' ou 'ambulances_sansoucy'
        
    Returns:
        str: Chemin du template créé
    """
    template_dir = os.path.join(settings.BASE_DIR, 'templates', 'word_templates')
    template_path = os.path.join(template_dir, f'contrat_{entity_name}_template.docx')
    
    # Si le template existe déjà, ne pas le recréer
    if os.path.exists(template_path):
        return template_path
    
    # Créer le dossier si nécessaire
    os.makedirs(template_dir, exist_ok=True)
    
    # Informations spécifiques à chaque entité
    entity_info = {
        'nantes_urgences': {
            'name': 'NANTES URGENCES SANSOUCY',
            'legal_form': 'SARL',
            'address': '12 Rue de la Chapelle\n44000 NANTES',
            'siret': '123 456 789 00012',
            'representative': 'M. Jean SANSOUCY, Gérant',
            'medical_service': 'Service de Santé au Travail de Loire-Atlantique',
        },
        'ambulances_sansoucy': {
            'name': 'AMBULANCES SANSOUCY',
            'legal_form': 'SAS',
            'address': '25 Avenue des Ambulances\n44300 NANTES',
            'siret': '987 654 321 00019',
            'representative': 'Mme. Marie SANSOUCY, Présidente',
            'medical_service': 'ASTIA - Service de Santé au Travail',
        }
    }
    
    info = entity_info.get(entity_name, entity_info['nantes_urgences'])
    
    # Créer le document
    doc = Document()
    
    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CONTRAT DE TRAVAIL')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # Sous-titre avec variables Jinja2
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('{{ contract_type_display }}')
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()
    
    # ENTRE LES SOUSSIGNÉS
    doc.add_paragraph('ENTRE LES SOUSSIGNÉS :').runs[0].bold = True
    doc.add_paragraph()
    
    # Employeur avec informations spécifiques
    p = doc.add_paragraph()
    p.add_run('L\'EMPLOYEUR\n').bold = True
    p.add_run(f'{info["name"]}\n')
    p.add_run(f'{info["legal_form"]}\n')
    p.add_run(f'{info["address"]}\n')
    p.add_run(f'SIRET : {info["siret"]}\n')
    p.add_run(f'Représentée par : {info["representative"]}\n')
    
    doc.add_paragraph()
    doc.add_paragraph('D\'UNE PART,').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Salarié avec variables Jinja2
    p = doc.add_paragraph()
    p.add_run('ET LE SALARIÉ\n').bold = True
    p.add_run('Nom : {{ employee_last_name }}\n')
    p.add_run('Prénom : {{ employee_first_name }}\n')
    p.add_run('Date de naissance : {{ employee_birth_date }}\n')
    p.add_run('Lieu de naissance : {{ employee_birth_place }}\n')
    p.add_run('N° Sécurité sociale : {{ employee_social_security }}\n')
    p.add_run('Adresse : {{ employee_address }}\n')
    
    doc.add_paragraph()
    doc.add_paragraph('D\'AUTRE PART,').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()
    
    # ARTICLE 1 - ENGAGEMENT
    doc.add_paragraph('ARTICLE 1 - ENGAGEMENT').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Le salarié est engagé à compter du {{ start_date }} en qualité de {{ employee_profession }}.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Type de contrat : {{ contract_type_display }}')
    
    # Conditions pour CDD
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('{% if end_date %}')
    p = doc.add_paragraph()
    p.add_run('Date de fin du contrat : {{ end_date }}')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    
    # ARTICLE 2 - PÉRIODE D'ESSAI
    doc.add_paragraph('ARTICLE 2 - PÉRIODE D\'ESSAI').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('{% if trial_end_date %}')
    p = doc.add_paragraph()
    p.add_run('Une période d\'essai est prévue jusqu\'au {{ trial_end_date }}.')
    p = doc.add_paragraph()
    p.add_run('{% else %}')
    p = doc.add_paragraph()
    p.add_run('Aucune période d\'essai n\'est prévue.')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    
    # ARTICLE 3 - DURÉE DU TRAVAIL
    doc.add_paragraph('ARTICLE 3 - DURÉE DU TRAVAIL').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'La durée hebdomadaire de travail est fixée à {{{{ working_hours_per_week }}}} heures.')
    
    doc.add_paragraph()
    
    # ARTICLE 4 - RÉMUNÉRATION
    doc.add_paragraph('ARTICLE 4 - RÉMUNÉRATION').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('{% if monthly_salary %}')
    p = doc.add_paragraph()
    p.add_run('Le salaire mensuel brut s\'élève à {{ monthly_salary }} euros.')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('{% if hourly_rate %}')
    p = doc.add_paragraph()
    p.add_run('Le taux horaire brut s\'élève à {{ hourly_rate }} euros.')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    
    # ARTICLE 5 - CONVENTION COLLECTIVE
    doc.add_paragraph('ARTICLE 5 - CONVENTION COLLECTIVE').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Le présent contrat est soumis à la {{ collective_agreement }}.')
    
    doc.add_paragraph()
    
    # ARTICLE 6 - MÉDECINE DU TRAVAIL
    doc.add_paragraph('ARTICLE 6 - MÉDECINE DU TRAVAIL').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'Le salarié relève du service de médecine du travail : {info["medical_service"]}')
    p = doc.add_paragraph()
    p.add_run('{% if occupational_health_service %}')
    p = doc.add_paragraph()
    p.add_run(' - {{ occupational_health_service }}')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Fait à {{ signature_place }}, le {{ signature_date }}')
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tableau de signatures
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    header_cells[0].text = 'L\'Employeur'
    header_cells[1].text = 'Le Salarié'
    
    # Espaces pour signatures
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Lignes de signature
    signature_cells = table.rows[1].cells
    for cell in signature_cells:
        cell.text = '\n\n\n'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sauvegarder
    doc.save(template_path)
    
    return template_path


def get_gender_agreements(gender):
    """
    Retourne les accords grammaticaux selon le genre
    
    Args:
        gender: 'M' (masculin) ou 'F' (féminin)
        
    Returns:
        dict: Dictionnaire avec les variantes grammaticales
    """
    if gender == 'F':
        return {
            'salarie': 'La salariée',
            'salarie_art': 'la salariée',
            'engage': 'engagée',
            'ne': 'née',
            'il_elle': 'elle',
            'son_sa': 'sa',
            'lui': 'elle',
        }
    else:
        return {
            'salarie': 'Le salarié',
            'salarie_art': 'le salarié',
            'engage': 'engagé',
            'ne': 'né',
            'il_elle': 'il',
            'son_sa': 'son',
            'lui': 'lui',
        }


def create_specific_template(entity_name, contract_type, gender):
    """
    Crée un template de contrat personnalisé pour une entité, type de contrat et genre spécifiques
    
    Args:
        entity_name: 'nantes_urgences' ou 'ambulances_sansoucy'
        contract_type: 'cdi' ou 'cdd'
        gender: 'M' (masculin) ou 'F' (féminin)
        
    Returns:
        str: Chemin du template créé
    """
    template_dir = os.path.join(settings.BASE_DIR, 'templates', 'word_templates')
    gender_suffix = 'homme' if gender == 'M' else 'femme'
    template_path = os.path.join(template_dir, f'contrat_{entity_name}_{contract_type}_{gender_suffix}.docx')
    
    # Si le template existe déjà, ne pas le recréer
    if os.path.exists(template_path):
        return template_path
    
    # Créer le dossier si nécessaire
    os.makedirs(template_dir, exist_ok=True)
    
    # Informations spécifiques à chaque entité
    entity_info = {
        'nantes_urgences': {
            'name': 'NANTES URGENCES SANSOUCY',
            'legal_form': 'SARL',
            'address': '12 Rue de la Chapelle\n44000 NANTES',
            'siret': '123 456 789 00012',
            'representative': 'M. Jean SANSOUCY, Gérant',
            'medical_service': 'Service de Santé au Travail de Loire-Atlantique',
        },
        'ambulances_sansoucy': {
            'name': 'AMBULANCES SANSOUCY',
            'legal_form': 'SAS',
            'address': '25 Avenue des Ambulances\n44300 NANTES',
            'siret': '987 654 321 00019',
            'representative': 'Mme. Marie SANSOUCY, Présidente',
            'medical_service': 'ASTIA - Service de Santé au Travail',
        }
    }
    
    info = entity_info.get(entity_name, entity_info['nantes_urgences'])
    agreements = get_gender_agreements(gender)
    
    # Créer le document
    doc = Document()
    
    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CONTRAT DE TRAVAIL')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # Sous-titre avec type de contrat
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contract_type_label = 'CDI - Contrat à Durée Indéterminée' if contract_type == 'cdi' else 'CDD - Contrat à Durée Déterminée'
    subtitle.add_run(contract_type_label)
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()
    
    # ENTRE LES SOUSSIGNÉS
    doc.add_paragraph('ENTRE LES SOUSSIGNÉS :').runs[0].bold = True
    doc.add_paragraph()
    
    # Employeur avec informations spécifiques
    p = doc.add_paragraph()
    p.add_run('L\'EMPLOYEUR\n').bold = True
    p.add_run(f'{info["name"]}\n')
    p.add_run(f'{info["legal_form"]}\n')
    p.add_run(f'{info["address"]}\n')
    p.add_run(f'SIRET : {info["siret"]}\n')
    p.add_run(f'Représentée par : {info["representative"]}\n')
    
    doc.add_paragraph()
    doc.add_paragraph('D\'UNE PART,').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Salarié avec variables Jinja2 et accords de genre
    p = doc.add_paragraph()
    p.add_run(f'ET {agreements["salarie"].upper()}\n').bold = True
    p.add_run('Nom : {{ employee_last_name }}\n')
    p.add_run('Prénom : {{ employee_first_name }}\n')
    p.add_run(f'Date de naissance : {{{{ employee_birth_date }}}} - {agreements["ne"]} à {{{{ employee_birth_place }}}}\n')
    p.add_run('N° Sécurité sociale : {{ employee_social_security }}\n')
    p.add_run('Adresse : {{ employee_address }}\n')
    
    doc.add_paragraph()
    doc.add_paragraph('D\'AUTRE PART,').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()
    
    # ARTICLE 1 - ENGAGEMENT
    doc.add_paragraph('ARTICLE 1 - ENGAGEMENT').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["salarie"]} est {agreements["engage"]}')
    
    if contract_type == 'cdd':
        p.add_run(' à compter du {{ start_date }} jusqu\'au {{ end_date }} en qualité de {{ employee_profession }}.')
    else:
        p.add_run(' à compter du {{ start_date }} en qualité de {{ employee_profession }}.')
    
    doc.add_paragraph()
    
    # ARTICLE 2 - PÉRIODE D'ESSAI
    doc.add_paragraph('ARTICLE 2 - PÉRIODE D\'ESSAI').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('{% if trial_end_date %}')
    p = doc.add_paragraph()
    p.add_run('Une période d\'essai est prévue jusqu\'au {{ trial_end_date }}.')
    p = doc.add_paragraph()
    p.add_run('{% else %}')
    p = doc.add_paragraph()
    p.add_run('Aucune période d\'essai n\'est prévue.')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    
    # ARTICLE 3 - DURÉE DU TRAVAIL
    doc.add_paragraph('ARTICLE 3 - DURÉE DU TRAVAIL').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["salarie"]} exercera {agreements["son_sa"]} activité à raison de {{{{ working_hours_per_week }}}} heures par semaine.')
    
    doc.add_paragraph()
    
    # ARTICLE 4 - RÉMUNÉRATION
    doc.add_paragraph('ARTICLE 4 - RÉMUNÉRATION').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('{% if monthly_salary %}')
    p = doc.add_paragraph()
    p.add_run(f'{agreements["salarie"]} percevra une rémunération mensuelle brute de {{{{ monthly_salary }}}} euros.')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    p = doc.add_paragraph()
    p.add_run('{% if hourly_rate %}')
    p = doc.add_paragraph()
    p.add_run('Le taux horaire brut s\'élève à {{ hourly_rate }} euros.')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    
    # ARTICLE 5 - CONVENTION COLLECTIVE
    doc.add_paragraph('ARTICLE 5 - CONVENTION COLLECTIVE').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Le présent contrat est soumis à la {{ collective_agreement }}.')
    
    doc.add_paragraph()
    
    # ARTICLE 6 - MÉDECINE DU TRAVAIL
    doc.add_paragraph('ARTICLE 6 - MÉDECINE DU TRAVAIL').runs[0].bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f'{agreements["salarie"]} relève du service de médecine du travail : {info["medical_service"]}')
    p = doc.add_paragraph()
    p.add_run('{% if occupational_health_service %}')
    p = doc.add_paragraph()
    p.add_run(' - {{ occupational_health_service }}')
    p = doc.add_paragraph()
    p.add_run('{% endif %}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Fait à {{ signature_place }}, le {{ signature_date }}')
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tableau de signatures
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    header_cells[0].text = 'L\'Employeur'
    header_cells[1].text = f'{agreements["salarie"]}'
    
    # Espaces pour signatures
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Lignes de signature
    signature_cells = table.rows[1].cells
    for cell in signature_cells:
        cell.text = '\n\n\n'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sauvegarder
    doc.save(template_path)
    
    return template_path






def generate_contract_document(contract):
    
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from contracts.cdd_templates_generator import get_gender_agreements, set_cell_background
    import re
    from datetime import datetime
    from io import BytesIO
    
    # Vérifier que c'est bien un contrat CDD
    if not contract.end_date:
        raise ValueError("Cette fonction ne génère que des contrats CDD (avec une date de fin)")
    
    employee = contract.employee
    gender = employee.gender if hasattr(employee, 'gender') and employee.gender else 'M'
    entity = contract.entity_template if contract.entity_template else 'nantes_urgences'
    contract_type = contract.contract_type
    
    # Obtenir les accords grammaticaux selon le genre
    agreements = get_gender_agreements(gender)
    
    # Créer le document
    doc = Document()
    
    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Déterminer les informations spécifiques à l'entité (TEXTE VERBATIM)
    if entity == 'nantes_urgences':
        entity_full_name = "NANTES URGENCES SANSOUCY"
        entity_siret = "48805076600028"
        entity_address = "8 Rue de Remouleur"
        entity_postal = "44800"
        entity_city = "SAINT-HERBLAIN"
        urssaf_code = "627201905366"
        representative_name = "Patrice BORÉ"
        representative_title = "Direction"
        location_work = "au sein des établissements de la société NANTES URGENCES SANSOUCY basée à St-Herblain ou Carquefou"
        
        if gender == 'M':
            authority_text = f"sera placé sous l'autorité hiérarchique de la direction de la société {entity_full_name}, il devra également respecter les missions et les directives qui lui seront données par les régulateurs."
        else:
            authority_text = f"sera placée sous l'autorité hiérarchique de la direction de la société {entity_full_name}, elle devra également respecter les missions et les directives qui lui seront données par les régulateurs."
        
        insurance_retirement = "Pour la retraite : UGRR - BP 501 - 75421 PARIS CEDEX 09."
        insurance_health_lines = []
        insurance_provident = "Pour la prévoyance : ALLIANZ PRÉVOYANCE - 1 cours Michelet - CS 30051 92076 PARIS LA DEFENSE CEDEX"
        
        employment_type = "DEA"
        hourly_rate = "13.06"
        monthly_salary = "1985.12"
        salary_base_date = "1er août 2025"
        city_signature = "St-Herblain"
        taxi_mention = ""
        planning_mention = "les plannings journaliers qui seront établis par les régulateurs"
        
        if gender == 'M':
            absence_text = f"Il s'oblige à prévenir sans délai la société {entity_full_name} de toute absence quelle qu'en soit la cause et de le justifier dans les 48h. Sinon cela correspondra à une absence injustifiée, qui, répétée, peut engendrer une sanction disciplinaire. "
        else:
            absence_text = f"Elle s'oblige à prévenir sans délai la société {entity_full_name} de toute absence quelle qu'en soit la cause et de le justifier dans les 48h. Sinon cela correspondra à une absence injustifiée, qui, répétée, peut engendrer une sanction disciplinaire. "
    
    else:  # ambulances_sansoucy
        entity_full_name = "SARL AMBULANCES SANSOUCY"
        entity_siret = "38026793000036"
        entity_address = "2 avenue de la Véra Cruz"
        entity_postal = "44600"
        entity_city = "SAINT NAZAIRE"
        urssaf_code = "527201905363"
        representative_name = "Bruno SANSOUCY"
        representative_title = "Gérant"
        location_work = f"à partir du lieu où se situe le siège de la société {entity_full_name}"
        
        if gender == 'M':
            authority_text = f"sera placé sous l'autorité hiérarchique des cogérants de la {entity_full_name}, il devra également respecter les plannings et les directives qui lui seront données par les régulateurs."
        else:
            authority_text = f"sera placée sous l'autorité hiérarchique des cogérants de la {entity_full_name}, elle devra également respecter les plannings et les directives qui lui seront données par les régulateurs."
        
        insurance_retirement = "Pour la retraite : UGRR - BP 501 - 75421 PARIS CEDEX 09."
        insurance_health_lines = ["Pour la complémentaire santé : HARMONIE MUTUELLE - 44824 ST HERBLAIN CEDEX"]
        insurance_provident = "Pour la prévoyance : ALLIANZ VIE - 1 Cours Michelet - 92076 PARIS LA DEFENSE CEDEX."
        
        employment_type = "AA"
        hourly_rate = "12.06"
        monthly_salary = "1833.12"
        salary_base_date = "1er janvier 2025"
        city_signature = "St-Nazaire"
        taxi_mention = f"Il est précisé qu'il pourra également être demandé à {agreements['civility']} {employee.user.first_name} {employee.user.last_name} de passer l'examen relatif à la conduite de taxi."
        planning_mention = "les plannings qui seront établis par les régulateurs"
        
        if gender == 'M':
            absence_text = f"Il s'oblige à prévenir sans délai la société {entity_full_name} de toute absence quelle qu'en soit la cause et, en cas de maladie, lui faire parvenir un arrêt de travail dans les 48 heures de l'arrêt. Sinon cela correspondra à une absence injustifiée, qui, répétée, peut engendrer une sanction disciplinaire. "
        else:
            absence_text = f"Elle s'oblige à prévenir sans délai la société {entity_full_name} de toute absence quelle qu'en soit la cause et, en cas de maladie, lui faire parvenir un arrêt de travail dans les 48 heures de l'arrêt. Sinon cela correspondra à une absence injustifiée, qui, répétée, peut engendrer une sanction disciplinaire. "
    
    # Formatter les dates
    def format_date_long(date):
        
        mois = ['janvier', 'février', 'mars', 'avril', 'mai', 'juin',
                'juillet', 'août', 'septembre', 'octobre', 'novembre', 'décembre']
        return f"{date.day:02d} {mois[date.month - 1]} {date.year}"
    
    def format_date_short(date):
        
        return date.strftime("%d/%m/%Y")
    
    start_date_long = format_date_long(contract.start_date)
    end_date_long = format_date_long(contract.end_date)
    birth_date_short = format_date_short(employee.birth_date)
    signature_date = format_date_short(contract.start_date)
    
    # Formater le numéro de sécurité sociale
    def format_ssn(ssn):
        
        ssn_clean = re.sub(r'\s+', '', str(ssn))
        if len(ssn_clean) == 15:
            return f"{ssn_clean[0]} {ssn_clean[1:3]} {ssn_clean[3:5]} {ssn_clean[5:7]} {ssn_clean[7:10]} {ssn_clean[10:13]} {ssn_clean[13:15]}"
        return ssn
    
    ssn_formatted = format_ssn(employee.social_security_number)
    
    # Déterminer le motif CDD
    if contract_type == 'accroissement_temporaire':
        contract_reason = "dans le cadre d'un accroissement temporaire de l'activité"
    else:
        contract_reason = "pour le remplacement d'un salarié absent"
    
    # Calculer la période d'essai (en jours)
    trial_period_days = contract.trial_period_days if hasattr(contract, 'trial_period_days') else 12
    
    # ==================== CONSTRUCTION DU DOCUMENT ====================
    
    # EN-TÊTE DU CONTRAT
    title = doc.add_paragraph()
    title_run = title.add_run("CONTRAT DE TRAVAIL À DURÉE DÉTERMINÉE")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # ENTRE
    p = doc.add_paragraph()
    p.add_run("ENTRE").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Informations de l'entreprise
    p = doc.add_paragraph()
    run = p.add_run(f"La société {entity_full_name}")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(f"SIRET {entity_siret}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(f"{entity_address}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(f"{entity_postal} {entity_city}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"URSSAF NANTES {urssaf_code}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # ET
    p = doc.add_paragraph("ET")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Informations de l'employé
    p = doc.add_paragraph()
    run = p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name}")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"{employee.address},").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"{employee.postal_code} {employee.city}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"Né{agreements['ne']} le : {birth_date_short} à {employee.birth_place}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"Immatricul é{agreements['ne']} : {ssn_formatted}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # IL A ETE RAPPELE CE QUI SUIT
    p = doc.add_paragraph()
    run = p.add_run("IL A ETE RAPPELE CE QUI SUIT")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Paragraphe d'embauche initial
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"est {agreements['engage']} en qualité d'ambulancier. ")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    
    # DPAE
    p = doc.add_paragraph()
    p.add_run("La déclaration préalable à l'embauche de ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("a été effectuée à l'URSSAF de Nantes et ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("pourra exercer auprès de cet organisme son droit d'accès et de rectification que lui confère la loi 78.17 du 6 janvier 1978.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # FONCTIONS
    p = doc.add_paragraph()
    run = p.add_run("FONCTIONS")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Paragraphe d'engagement
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"est {agreements['engage']} à durée déterminée à partir du {start_date_long} jusqu'au {end_date_long}, avec une période d'essai de {trial_period_days} jours. ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"est {agreements['engage']} à durée déterminée, {contract_reason}. ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Fonctions à effectuer
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("effectuera notamment au sein de notre entreprise les fonctions définies à l'annexe 1 du présent contrat (dont copie ci-jointe).")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # ARTICLE 1 : ATTRIBUTION ET EMPLOI
    p = doc.add_paragraph()
    run = p.add_run("ARTICLE 1 : ATTRIBUTION ET EMPLOI")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"est {agreements['engage']} par la société {entity_full_name} en qualité d'ambulancier, de la convention collective des transports routiers (IDCC 16), applicable à l'activité.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Missions
    p = doc.add_paragraph()
    p.add_run("En cette qualité ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("aura pour mission :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("Le transport de personnes en ambulance et VSL")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("Et plus généralement le transport de personnes.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Mention taxi pour Ambulances Sansoucy
    if taxi_mention:
        p = doc.add_paragraph(taxi_mention)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
    
    # Autorité hiérarchique
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(authority_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Espaces avant Article 2
    for _ in range(8):
        doc.add_paragraph()
    
    # ARTICLE 2 : CONDITIONS GÉNÉRALES DE TRAVAIL
    p = doc.add_paragraph()
    run = p.add_run("ARTICLE 2 : CONDITIONS GÉNÉRALES DE TRAVAIL")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # 2.1 Lieu de travail
    p = doc.add_paragraph()
    run = p.add_run("2.1 Lieu de travail")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"exercera ses fonctions {location_work}.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Toutefois, ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    if gender == 'M':
        p.add_run(f"{agreements['accepte']} par avance qu'en fonction des nécessités de l'entreprise, il soit amené à changer de lieu de travail, et ce, dans les zones géographiques où la société exerce ou exercera son activité. ")
    else:
        p.add_run(f"{agreements['accepte']} par avance qu'en fonction des nécessités de l'entreprise, qu'elle soit amenée à changer de lieu de travail, et ce, dans les zones géographiques où la société exerce ou exercera son activité.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # 2.2 Congés payés
    p = doc.add_paragraph()
    run = p.add_run("2.2 Congés payés")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    if gender == 'M':
        p.add_run("Conformément aux dispositions légales en vigueur et celles de la convention collective applicable, le salarié bénéficiera de 2.08 jours ouvrés de congés par mois de travail effectif, acquis sur la période courant du 1er juin au 31 mai de l'année suivante.")
    else:
        p.add_run("Conformément aux dispositions légales en vigueur et celles de la convention collective applicable, la salariée bénéficiera de 2.08 jours ouvrés de congés par mois de travail effectif, acquis sur la période courant du 1er juin au 31 mai de l'année suivante.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("La direction fixe les périodes de congés, en fonction des nécessités du service, après concertation du personnel.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # 2.3 Caisse de retraite complémentaire
    p = doc.add_paragraph()
    run = p.add_run("2.3 Caisse de retraite complémentaire")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"sera {agreements['affilie']} :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph(insurance_retirement)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for health_line in insurance_health_lines:
        p = doc.add_paragraph(health_line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph(insurance_provident)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # 2.4 Conditions générales d'exercice des fonctions
    p = doc.add_paragraph()
    run = p.add_run("2.4 Conditions générales d'exercice des fonctions")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"s'engage à se conformer strictement aux instructions de la direction concernant les conditions d'exécution du travail, et à respecter {planning_mention}.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Suite des paragraphes 2.4 concernant le permis
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("declare ne faire l'objet d'aucune restriction administrative ou judiciaire quant à l'utilisation de son permis de conduire, toutes catégories confondues.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("s'engage également à produire chaque année, à la date anniversaire de son contrat de travail, son permis de conduire, de même qu'à toute époque de l'année, sur simple demande de la direction.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("s'engage à avertir immédiatement son employeur de toute suspension de permis ou de tout événement susceptible de remettre en cause sa faculté de conduire les véhicules de l'entreprise.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Il est également rappelé à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("que toute suspension ou invalidation de son permis de conduire sera susceptible d'entraîner la rupture de son contrat de travail.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Il est enfin rappelé à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"que les véhicules appartenant à la société {entity_full_name} ne peuvent, sauf accord écrit de la direction, être utilisés à des fins personnelles ou servir à transporter des personnes étrangères à la société (hormis les clients).")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Il est également rappelé à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    if gender == 'M':
        p.add_run("qu'il est tenu de nettoyer et de laisser propre les locaux mis à sa disposition.")
    else:
        p.add_run("qu'elle est tenue de nettoyer et de laisser propre les locaux mis à sa disposition.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Amendes
    p = doc.add_paragraph()
    if entity == 'nantes_urgences':
        p.add_run(f"Toutes les amendes encourues lors de l'utilisation du véhicule de la société {entity_full_name} seront dues par le conducteur mis en cause et donneront lieu au remplissage du formulaire de requête en exonération. Si toutefois la société devait faire l'avance des sommes dues cela donnerait lieu à un remboursement par le salarié concerné. De plus, en cas d'infraction routière, l'entreprise, en application de l'article L121-6 du code de la route devra divulguer l'identité du salarié ayant commis l'infraction aux autorités compétentes. ")
    else:
        p.add_run(f"Toutes les amendes encourues lors de l'utilisation du véhicule de la {entity_full_name} seront dues par le conducteur mis en cause et donneront lieu au remplissage du formulaire de requête en exonération. Si toutefois la société devait faire l'avance des sommes dues cela donnerait lieu à un remboursement par le salarié concerné. De plus, en cas d'infraction routière, l'entreprise, en application de l'article L121-6 du code de la route devra divulguer l'identité du salarié ayant commis l'infraction aux autorités compétentes. ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # ARTICLE 3 : DISCIPLINE ET SÉCURITÉ
    p = doc.add_paragraph()
    run = p.add_run("ARTICLE 3 : DISCIPLINE ET SÉCURITÉ")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("déclare que toutes les coordonnées figurant à l'entête des présentes sont exactes, et s'oblige à prévenir l'employeur de toutes modifications les affectant.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph(absence_text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p = doc.add_paragraph("En cas de maladie, il faut faire parvenir un arrêt de travail dans les 48 heures de l'arrêt au service ressources humaines de l'entreprise. ")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    if gender == 'M':
        p.add_run("Le salarié s'engage à ne pas entrer au service d'autres employeurs pour une durée du travail qui, ajoutée à celle convenue dans le présent contrat, entraînerait un dépassement des durées maximales légales telles que prévues par le Code du travail.")
    else:
        p.add_run("La salariée s'engage à ne pas entrer au service d'autres employeurs pour une durée du travail qui, ajoutée à celle convenue dans le présent contrat, entraînerait un dépassement des durées maximales légales telles que prévues par le Code du travail.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("reconnaît avoir pris connaissance du règlement intérieur en vigueur dans l'établissement. Tout manquement au présent règlement pourra donner lieu à des poursuites disciplinaires et à un éventuel licenciement pour faute. ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # ARTICLE 4 : CONFIDENTIALITE ET SECRET PROFESSIONNEL
    p = doc.add_paragraph()
    run = p.add_run("ARTICLE 4 : CONFIDENTIALITE ET SECRET PROFESSIONNEL")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Compte tenu des fonctions confiées à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name}, ").bold = True
    p.add_run("celui-ci est tenu par un secret professionnel tant en ce qui concerne l'identité des clients transportés que leur destination.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ARTICLE 5 : SALAIRE
    p = doc.add_paragraph()
    run = p.add_run("ARTICLE 5 : SALAIRE")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Le salaire de ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name}, ").bold = True
    p.add_run("SMPG (Salaire Minimum Professionnel Garanti), se décompose comme suit :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Table de salaire
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # En-tête
    if entity == 'nantes_urgences':
        table.cell(0, 0).text = "Emploi"
        table.cell(0, 1).text = "Taux"
        table.cell(0, 2).text = f"A compter du {salary_base_date} (base 152h / mois)"
    else:
        table.cell(0, 0).text = "Dates         \nEmplois"
        table.cell(0, 1).text = "Taux Horaire"
        table.cell(0, 2).text = f"A compter du {salary_base_date} (base 152h / mois)"
    
    # Données
    table.cell(1, 0).text = employment_type
    table.cell(1, 1).text = hourly_rate
    table.cell(1, 2).text = f"{monthly_salary} €"
    
    # Mise en forme
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    # Fond grisé
    for i in range(3):
        set_cell_background(table.rows[0].cells[i], 'D9D9D9')
    set_cell_background(table.rows[1].cells[0], 'D9D9D9')
    
    doc.add_paragraph()
    
    # Indemnités
    p = doc.add_paragraph()
    p.add_run("En sus du SMPG, il pourra être versé à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("les indemnités suivantes, dans les termes de l'accord du 04 mai 2000 :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("IDAJ (Indemnité Dépassement d'amplitude Journalière)")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("Tâches complémentaires : elles se trouvent définies à l'accord cadre et donneront lieu à un paiement spécifique chaque fois qu'elles auront été effectuées.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Personnel ambulancier :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Type 1  - Conduite de tous véhicules non sanitaires de moins de 10 places")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("            - Transport de corps avant mise en bière")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("            - Transport, livraison, installation et entretien de matériel médical")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Type 2  - Funéraire, tâches d'exécution (porteurs, etc.) places")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("     - Taxi (titulaire du certificat de taxi)")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Type 3  - Régulation")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("            - Autre activité funéraire")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("     - Mécanique, réparation automobile")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Rémunération
    p = doc.add_paragraph()
    p.add_run("Rémunération :").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("L'accomplissement de ces tâches complémentaires donnera lieu, sur le mois considéré où elles ont été effectuées, à une majoration des montants du SMPG dans les conditions suivantes :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Personnel ambulancier :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("Type 1  2%")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("Type 2  5%")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("Type 3  10%")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Ces majorations ne se cumulent pas, seule la plus forte est appliquée en cas d'exercice de plusieurs tâches. Elles apparaîtront sur une ligne à part sur le bulletin de salaire. \nIl peut à tout moment être sollicité du salarié l'exécution desdites tâches complémentaires. ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # Indemnité d'ancienneté
    p = doc.add_paragraph()
    p.add_run("Enfin, il sera versé à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(":")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("Une indemnité d'ancienneté qui sera calculée selon les modalités suivantes, ressortant à l'accord cadre du 04 mai.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Personnel ouvrier :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("2% après 2 ans")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("4% après 5 ans ininterrompus dans l'entreprise")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("6% après 10 ans")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("8% après 15 ans.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    # ARTICLE 6 : PERMANENCES
    p = doc.add_paragraph()
    p.add_run("ARTICLE 6 : PERMANENCES").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Il pourra être demandé à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("d'effectuer des permanences dans les lieux désignés par l'entreprise.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("Celles-ci correspondent aux gardes de  ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("              - Nuits")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("            - Samedis")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("            - Dimanches et jours fériés.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ARTICLE 7 : HEURES SUPPLEMENTAIRES
    p = doc.add_paragraph()
    p.add_run("ARTICLE 7 : HEURES SUPPLEMENTAIRES").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run(f"devra tenir des feuilles de route sur lesquelles {agreements['il_elle']} fera figurer les éléments nécessaires à l'établissement de {agreements['son_sa']} temps de travail. Ces feuilles de route lui seront fournies, en support, par l'entreprise. Elles devront être remplies de façon quotidienne et seront visées par l'entreprise.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph()
    p.add_run("Ces feuilles de route seront à dispositions du salarié, afin de permettre à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("d'effectuer tout contrôle qui lui paraîtra utile sur l'établissement du salaire.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Il pourra être demandé à ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("d'effectuer des heures supplémentaires. Elles seront décomptées et rémunérées dans les termes de l'accord du 04 mai 2000, et en toute hypothèse, dans les termes des dispositions légales et réglementaires régissant notre profession. ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Elles pourront faire l'objet :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("Soit d'un paiement majoré ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("Soit d'un repos compensateur de remplacement")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph("Soit les deux à la fois.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    p = doc.add_paragraph("Les heures supplémentaires ayant donné lieu à un repos compensateur de remplacement ne s'imputent pas sur le contingent annuel d'heures supplémentaires.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    for _ in range(4):
        doc.add_paragraph()
    
    # ARTICLE 8 : INDEMNITÉ DE FIN DE CONTRAT
    p = doc.add_paragraph()
    p.add_run("ARTICLE 8 : INDEMNITÉ DE FIN DE CONTRAT").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph()
    p.add_run("Lorsque, à l'issue d'un contrat de travail à durée déterminée, les relations contractuelles de travail ne se poursuivent pas par un contrat à durée indéterminée, ")
    p.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} ").bold = True
    p.add_run("a droit, à titre de complément de salaire, à une indemnité de fin de contrat destinée à compenser la précarité de sa situation.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("Cette indemnité est égale à 10 % de la rémunération totale brute versée au salarié.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("L'indemnité de fin de contrat n'est pas due :")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("1° Lorsque le contrat est conclu au titre du 3° de l'article L. 1242-2 ou de l'article L. 1242-3, sauf dispositions conventionnelles plus favorables ;")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("2° Lorsque le contrat est conclu avec un jeune pour une période comprise dans ses vacances scolaires ou universitaires ;")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph()
    if gender == 'M':
        p.add_run("3° Lorsque le salarié refuse d'accepter la conclusion d'un contrat de travail à durée indéterminée pour occuper le même emploi ou un emploi similaire, assorti d'une rémunération au moins équivalente ;")
    else:
        p.add_run("3° Lorsque la salariée refuse d'accepter la conclusion d'un contrat de travail à durée indéterminée pour occuper le même emploi ou un emploi similaire, assorti d'une rémunération au moins équivalente ;")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph()
    if gender == 'M':
        p.add_run("4° En cas de rupture anticipée du contrat due à l'initiative du salarié, à sa faute grave ou à un cas de force majeure.")
    else:
        p.add_run("4° En cas de rupture anticipée du contrat due à l'initiative de la salariée, à sa faute grave ou à un cas de force majeure.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # ARTICLE 9 : RENOUVELLEMENT
    p = doc.add_paragraph()
    p.add_run("ARTICLE 9 : RENOUVELLEMENT").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p = doc.add_paragraph("Le présent contrat pourra faire l'objet de deux renouvellements formalisés par un accord entre les parties.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    
    # Fait à
    p = doc.add_paragraph()
    p.add_run(f"Fait à {city_signature}, le {signature_date} : ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table de signatures
    sig_table = doc.add_table(rows=2, cols=2)
    
    # Première ligne - titres
    cell_employee = sig_table.cell(0, 0)
    p_emp = cell_employee.paragraphs[0]
    p_emp.add_run(f"{agreements['civility']} {employee.user.first_name} {employee.user.last_name} :").bold = True
    cell_employee.add_paragraph('Mention "Bon pour accord, lu et approuvé" ')
    
    cell_company = sig_table.cell(0, 1)
    p_comp = cell_company.paragraphs[0]
    run = p_comp.add_run(f"     {entity_full_name}")
    run.bold = True
    p_comp = cell_company.add_paragraph(f"     Monsieur {representative_name}")
    p_comp.runs[0].bold = True
    p_comp = cell_company.add_paragraph(f"     {representative_title}")
    p_comp.runs[0].bold = True
    
    # Deuxième ligne - espaces pour signatures
    sig_table.cell(1, 0).text = ""
    sig_table.cell(1, 1).text = ""
    
    # Sauvegarder dans un buffer BytesIO
    file_buffer = BytesIO()
    doc.save(file_buffer)
    file_buffer.seek(0)
    
    # Générer le nom du fichier
    start_date_filename = format_date_short(contract.start_date).replace('/', '-')
    filename = f"CDD_{employee.user.last_name}_{employee.user.first_name}_{entity}_{start_date_filename}.docx"
    
    return filename, file_buffer.getvalue()

